"""
Genera Memoria.docx con el formato exigido por la normativa ASIR:
- Arial 12pt, 1.5 interlineado, justificado
- Márgenes: Sup 3cm, Inf 2.5cm, Izq 3cm, Der 2.5cm
- Capítulos: 14pt, negrita, mayúsculas
- Subcapítulos: 12pt, negrita
- Numeración de páginas en pie
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE = r"D:\PROYECTOS\TFG"
CHAPTERS = [
    r"Memoria\1 - Estudio del problema y análisis del sistema.md",
    r"Memoria\2 - Recursos necesarios.md",
    r"Memoria\3 - Implementación y documentación técnica.md",
    r"Memoria\4 - Fase de pruebas.md",
    r"Memoria\5 - Conclusiones finales.md",
    r"Memoria\6 - Documentación del sistema desarrollado.md",
    r"Memoria\7 - Bibliografía.md",
]
OUT = os.path.join(BASE, "Entregables", "Memoria-TFG-Liher-Rios.docx")

def set_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(3)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

def add_page_numbers(doc):
    """Añade número de página en el pie centrado."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.name = 'Arial'
        run.font.size = Pt(10)

def set_para_format(para, size=12, bold=False, align='justify', space_before=0, space_after=6):
    para.alignment = {
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'center':  WD_ALIGN_PARAGRAPH.CENTER,
        'left':    WD_ALIGN_PARAGRAPH.LEFT,
    }[align]
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    from docx.shared import Pt as PT
    from docx.oxml.ns import qn as QN
    # 1.5 line spacing
    pf.line_spacing_rule = 1  # EXACTLY
    from docx.shared import Pt as _Pt
    pf.line_spacing = _Pt(18)  # 12pt * 1.5

def apply_run_format(run, size=12, bold=False, italic=False):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.bold  = bold
    run.italic = italic

def add_cover(doc):
    """Portada conforme a normativa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("SALESIANOS LOS BOSCOS")
    apply_run_format(r, 14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("CFGS Administración de Sistemas Informáticos en Red (ASIR) — Dual")
    apply_run_format(r, 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    r = p.add_run("IMPLEMENTACIÓN DE AUTOMATIZACIONES\nPARA ENTORNOS EDUCATIVOS")
    apply_run_format(r, 16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Proyecto de Fin de Ciclo")
    apply_run_format(r, 12, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("Alumno: Liher Ríos Ruiz")
    apply_run_format(r, 12, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Curso académico: 2024-2025")
    apply_run_format(r, 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tutor: *** RELLENAR NOMBRE DEL TUTOR ***")
    apply_run_format(r, 12, bold=True)
    r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # rojo visible

    doc.add_page_break()

def parse_inline(text):
    """Devuelve lista de (texto, bold, italic, code)."""
    segments = []
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            segments.append((text[last:m.start()], False, False, False))
        if m.group(2):
            segments.append((m.group(2), True, True, False))
        elif m.group(3):
            segments.append((m.group(3), True, False, False))
        elif m.group(4):
            segments.append((m.group(4), False, True, False))
        elif m.group(5):
            segments.append((m.group(5), False, False, True))
        last = m.end()
    if last < len(text):
        segments.append((text[last:], False, False, False))
    return segments

def add_inline(para, text, base_size=12):
    for seg_text, bold, italic, code in parse_inline(text):
        r = para.add_run(seg_text)
        r.font.name = 'Courier New' if code else 'Arial'
        r.font.size = Pt(11 if code else base_size)
        r.bold   = bold
        r.italic = italic

def process_table(doc, lines, start_idx):
    """Procesa una tabla markdown. Devuelve el índice final."""
    table_lines = []
    i = start_idx
    while i < len(lines) and '|' in lines[i]:
        table_lines.append(lines[i])
        i += 1

    # Filtrar separadores (|---|)
    rows = [l for l in table_lines if not re.match(r'^\|[\s\-:|]+\|', l)]
    if not rows:
        return i

    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cells)

    if not parsed:
        return i

    cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=cols)
    table.style = 'Table Grid'

    for ri, row in enumerate(parsed):
        for ci, cell_text in enumerate(row):
            if ci >= cols:
                break
            cell = table.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            r = p.add_run(clean)
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            if ri == 0:
                r.bold = True
            p.paragraph_format.line_spacing = Pt(14)

    doc.add_paragraph()
    return i

def process_chapter(doc, filepath):
    with open(filepath, encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Bloque de código
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(6)
                p.paragraph_format.line_spacing  = Pt(14)
                r = p.add_run('\n'.join(code_lines))
                r.font.name = 'Courier New'
                r.font.size = Pt(9)
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tabla
        if '|' in line and line.strip().startswith('|'):
            i = process_table(doc, lines, i)
            continue

        # Headings
        h1 = re.match(r'^# (.+)', line)
        h2 = re.match(r'^## (.+)', line)
        h3 = re.match(r'^### (.+)', line)
        h4 = re.match(r'^#### (.+)', line)

        if h1:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after  = Pt(12)
            p.paragraph_format.line_spacing  = Pt(21)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(h1.group(1).upper())
            r.font.name = 'Arial'
            r.font.size = Pt(14)
            r.bold = True
            i += 1
            continue

        if h2:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.line_spacing  = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(h2.group(1))
            r.font.name = 'Arial'
            r.font.size = Pt(12)
            r.bold = True
            i += 1
            continue

        if h3:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.line_spacing  = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(h3.group(1))
            r.font.name = 'Arial'
            r.font.size = Pt(12)
            r.bold = True
            i += 1
            continue

        if h4:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.line_spacing  = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(h4.group(1))
            r.font.name = 'Arial'
            r.font.size = Pt(12)
            r.bold = True
            r.italic = True
            i += 1
            continue

        # Lista numerada
        num_match = re.match(r'^(\d+)\. (.+)', line)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.line_spacing  = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, num_match.group(2))
            i += 1
            continue

        # Lista con guión/asterisco
        bullet_match = re.match(r'^[-*] (.+)', line)
        if bullet_match:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.line_spacing  = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, bullet_match.group(1))
            i += 1
            continue

        # Línea vacía
        if not line.strip():
            i += 1
            continue

        # Párrafo normal
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.line_spacing  = Pt(18)
        add_inline(p, line)
        i += 1


def main():
    doc = Document()

    # Eliminar estilo por defecto y configurar
    set_margins(doc)

    # Estilo Normal base
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    add_cover(doc)
    add_page_numbers(doc)

    for idx, rel_path in enumerate(CHAPTERS):
        path = os.path.join(BASE, rel_path)
        if not os.path.exists(path):
            print(f"[SKIP] No encontrado: {path}")
            continue
        print(f"[OK] Procesando: {rel_path}")
        process_chapter(doc, path)
        if idx < len(CHAPTERS) - 1:
            doc.add_page_break()

    doc.save(OUT)
    print(f"\nGenerado: {OUT}")

if __name__ == '__main__':
    main()
